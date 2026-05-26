from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


INPUT_PATH = Path("/Users/chartat1/Documents/Papers/SuperDARN_winds/paper_first_submission_AnGeo.docx")
OUTPUT_PATH = Path(
    "/Users/chartat1/Documents/Papers/SuperDARN_winds/paper_first_submission_AnGeo_submissionformatted.docx"
)


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def clear_runs(paragraph: Paragraph) -> None:
    for run in paragraph.runs[::-1]:
        paragraph._p.remove(run._r)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    clear_runs(paragraph)
    paragraph.add_run(text)


def remove_numbering(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def qn(tag: str) -> str:
    prefix, tagroot = tag.split(":")
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    return f"{{{nsmap[prefix]}}}{tagroot}"


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def paragraph_index(doc: Document, target: Paragraph) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph._p is target._p:
            return idx
    raise ValueError("Paragraph not found in document.")


def main() -> None:
    doc = Document(INPUT_PATH)

    # Add corresponding author line after affiliations if one is not already present.
    if not any("Correspondence to:" in p.text for p in doc.paragraphs[:12]):
        affil_para = next(
            p for p in doc.paragraphs if p.text.strip() == "4 Pennsylvania State University, University Park, PA, USA"
        )
        corr = insert_paragraph_after(affil_para, "Correspondence to: Alex Chartier (alex.chartier@jhuapl.edu)")
        corr.style = affil_para.style

    # Remove the non-AnGeo front-matter extras between the abstract and Introduction.
    start_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Plain text summary")
    end_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Introduction")
    for idx in range(end_idx - 1, start_idx - 1, -1):
        delete_paragraph(doc.paragraphs[idx])

    # Refresh paragraph list after deletions.
    doc = Document(INPUT_PATH) if False else doc

    # Normalize section headings.
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "1.6 The JAGUAR-DAS Whole Neutral Atmosphere Reanalysis":
            set_paragraph_text(paragraph, "The JAGUAR-DAS Whole Neutral Atmosphere Reanalysis")
            paragraph.style = "Heading 2"
        elif text == "Method":
            set_paragraph_text(paragraph, "Methods")
            paragraph.style = "Heading 1"
        elif text.startswith("The observed and modeled counts resulted in similar zonal wind distributions."):
            paragraph.style = "Normal"
        elif text.startswith(
            "The meteor count model was applied in comparing SuperDARN wind observations against those observed"
        ):
            paragraph.style = "Normal"
        elif text == "Data Availability Statement":
            set_paragraph_text(paragraph, "Code and data availability")
            paragraph.style = "Heading 1"
            remove_numbering(paragraph)

    # Update end matter to align with Copernicus sectioning.
    data_heading = next(p for p in doc.paragraphs if p.text.strip() == "Code and data availability")
    data_para = next(
        p
        for p in doc.paragraphs[
            paragraph_index(doc, data_heading) + 1 :
        ]
        if p.text.strip()
    )
    ack_heading = next(p for p in doc.paragraphs if p.text.strip() == "Acknowledgements")
    ack_text = next(
        p
        for p in doc.paragraphs[
            paragraph_index(doc, ack_heading) + 1 :
        ]
        if p.text.strip()
    )
    # Insert Author contributions and Competing interests after code/data availability.
    author_heading = insert_paragraph_after(data_para, "Author contributions", "Heading 1")
    author_text = insert_paragraph_after(
        author_heading,
        (
            "ATC conceived the study, developed the software, processed the data, carried out the analysis, "
            "prepared the figures, and drafted the manuscript. RP and RLM contributed to software development, "
            "data processing, and manuscript editing. DJ, JC, TR, RL, and WB contributed data, scientific input, "
            "and manuscript editing. All authors reviewed and approved the final manuscript."
        ),
        "Normal",
    )
    comp_heading = insert_paragraph_after(author_text, "Competing interests", "Heading 1")
    comp_text = insert_paragraph_after(
        comp_heading,
        "The contact author has declared that none of the authors has any competing interests.",
        "Normal",
    )

    # Replace acknowledgement text and move funding to a dedicated section.
    set_paragraph_text(
        ack_text,
        "The authors thank the SuperDARN and meteor-radar teams and the JAWARA data providers for making this study possible.",
    )
    ack_text.style = "Normal"

    fin_heading = insert_paragraph_after(ack_text, "Financial support", "Heading 1")
    fin_text = insert_paragraph_after(
        fin_heading,
        "ATC, RLM and RP acknowledge support from NASA 80NSSC23K0094. ATC acknowledges support from NSF 1934973 and 2426201.",
        "Normal",
    )

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
